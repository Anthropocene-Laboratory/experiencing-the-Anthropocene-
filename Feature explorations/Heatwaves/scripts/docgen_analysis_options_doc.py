from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OUT = ROOT / "Feature explorations" / "Heatwaves" / "Analysis_options_heatwave_population_europe.docx"

# Design preset: decision_memo (standard_business_brief).
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
MUTED = "666666"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_label=None):
    p = doc.add_paragraph(style="Normal")
    if bold_label:
        r = p.add_run(bold_label)
        set_run_font(r, bold=True)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GREY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=11, color=INK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_comparison_table(doc):
    widths = [1720, 1650, 3010, 2980]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, widths)
    headers = ["Approach", "Best use", "Strengths", "Constraints"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=INK, bold=True)

    rows = [
        (
            "QGIS only",
            "Visual exploration, manual checks and first map prototypes.",
            "Fast to inspect layers, projections and temporal behaviour; no coding required for simple overlays.",
            "Manual processing is hard to reproduce at 30 years; limited for repeated raster calculations and sensitivity tests.",
        ),
        (
            "R spatial workflow",
            "Reproducible raster calculations, annual summaries and statistical analysis.",
            "Transparent scripts; terra for percentile climatology, exactextractr for count-conserving aggregation and zonal summaries, sf for country joins.",
            "Requires package setup and debugging; a map interface is not the natural final product.",
        ),
        (
            "R + QGIS",
            "Recommended initial workflow: calculate in R, validate visually in QGIS.",
            "Combines reproducibility with rapid spatial quality control; keeps analytical and display steps distinct.",
            "Uses two environments and therefore needs clear file naming and metadata discipline.",
        ),
        (
            "R + QGIS + Shiny",
            "Published interactive map after outputs have been validated.",
            "A year slider, layer controls and country pop-ups can be shared without distributing GIS software.",
            "Adds deployment, performance and maintenance work; should not be the first place to discover data errors.",
        ),
        (
            "Python / xarray (+ CDO)",
            "Large-scale retrieval and array processing of the NetCDF climate inputs (E-OBS and ERA5-Land).",
            "Natural fit for CDS access and gridded daily data; xarray/CDO are efficient for percentile climatologies and batch processing.",
            "A second runtime to install and keep reproducible, and a second place to maintain later. The human learning cost is largely removed when the code is AI-written, so this is a minor constraint if R stays the main environment.",
        ),
    ]
    for row_values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_source_table(doc):
    widths = [1500, 1450, 3220, 3190]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, widths)
    headers = ["Input", "Recommended source", "Why this source", "Trade-off / fallback"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=INK, bold=True)

    rows = [
        (
            "Temperature (primary)",
            "E-OBS (CDS insitu-gridded-observations-europe)",
            "Station-based, native 0.1° regular grid, land-only, 1950-present. Outperforms reanalysis where station density is high and avoids the ERA5-Land coastal land-sea mixing bias that depresses temperatures exactly where coastal populations concentrate. Daily TX, TN and HU all in one product. Sources: Cornes et al. 2018 (E-OBS ensemble); Muñoz-Sabater et al. 2021 (ERA5-Land); coastal cool bias reported in heat-exposure assessments (e.g. Lazio two-stage heat-mortality study, 2025).",
            "Accuracy degrades in sparse-station regions (parts of Eastern Europe, high mountains). Treat those areas with caution and use the ensemble spread as an uncertainty signal.",
        ),
        (
            "Temperature (cross-check)",
            "ERA5-Land daily statistics",
            "Physically consistent global reanalysis at 0.1°, complete in space, useful to confirm E-OBS in data-sparse areas and as a robustness comparison for the heatwave classification.",
            "Coastal cool bias and smoother extremes; do not use as the sole source for a population-exposure metric without the E-OBS comparison.",
        ),
        (
            "Population (primary)",
            "GHS-POP R2023A, 5-year epochs",
            "Consistent dasymetric population counts back to 1975, the only product that covers the full 1991-2020 window. Linearly interpolated to annual layers between adjacent epochs.",
            "5-year epochs only; linear interpolation is an explicit assumption that does not recover within-interval relocation.",
        ),
        (
            "Population (robustness)",
            "WorldPop annual",
            "Natively annual at ~100 m, useful to check the GHS-POP interpolation from 2000 onward.",
            "Starts in 2000, so it cannot cover 1991-1999; different method, so use for sensitivity only, not as the primary series.",
        ),
        (
            "Age / life stage",
            "Eurostat demo_pjan (65+, 75+ shares)",
            "National annual age structure, joined to country-year exposure summaries after spatial aggregation.",
            "National, not gridded; do not impute a uniform age raster. Optional NUTS3 detail (demo_r_pjangrp3) only for 2014-2020.",
        ),
    ]
    for row_values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_stage_table(doc):
    widths = [1700, 2080, 5580]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, ["Stage", "Primary tool", "Deliverable and decision rule"]):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10, color=INK, bold=True)
    rows = [
        ("Acquire", "CDS API + direct download", "E-OBS daily TX/TN (primary) and ERA5-Land daily statistics (cross-check), GHS-POP epochs, GISCO boundaries and Eurostat age tables. Preserve source files, request JSON and checksums."),
        ("Prepare", "R", "Crop the European study area, linearly interpolate annual GHS-POP, then aggregate population to the 0.1° grid by count-conserving area-weighted sum (exactextractr). Retain the raw-to-processed lineage."),
        ("Calculate", "R", "Local calendar-day P90 thresholds for TX and TN, two-or-more-day compound heatwave flag, intensity and annual exposed person-days. Keep each annual component as a separate output."),
        ("Validate", "QGIS", "Check CRS, extent, no-data, population conservation, E-OBS vs ERA5-Land agreement, selected heatwave years and country totals before publication."),
        ("Share", "Shiny", "Interactive annual person-day map with a fixed colour scale and year slider, served as COG / pre-rendered tiles. Build only after validation."),
    ]
    for values in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("EXPERIENCING THE ANTHROPOCENE  |  PHASE 3")
    set_run_font(r, size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    r = footer.add_run("Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(footer)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Analysis options for European heatwave population exposure")
    set_run_font(r, size=22, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Decision note for the 1991-2020 interactive mapping prototype")
    set_run_font(r, size=12, color=MUTED, italic=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    for label, value in [("Scope", "European study area"), ("Primary output", "annual heatwave person-days exposed"), ("Status", "recommended workflow for initial implementation")]:
        r = meta.add_run(f"{label}: ")
        set_run_font(r, size=10, bold=True, color=INK)
        r = meta.add_run(value + "    ")
        set_run_font(r, size=10, color=MUTED)

    add_callout(
        doc,
        "Recommended decision",
        "Use E-OBS (station-based, land-only, native 0.1°) as the primary climate input, with ERA5-Land retained only as a sensitivity cross-check. Calculate the reproducible exposure layers in R, conserving population counts through area-weighted summation (exactextractr), validate in QGIS, and build a Shiny map only after the layers pass quality checks.",
    )

    add_heading(doc, "1. Analytical target", 1)
    add_body(doc, "The first map should not attempt to display every relevant variable at once. It should map annual heatwave person-days exposed: a Layer A thermal event filtered by Layer B population distribution.")
    add_body(doc, "Age is retained as a separate Layer B country-year analysis, not imposed on every grid cell. This avoids assuming that a national age structure is spatially uniform within a country.")

    add_heading(doc, "2. Workflow options", 1)
    add_body(doc, "The options below are complementary rather than mutually exclusive. The choice is about where each job is best performed.")
    add_comparison_table(doc)

    add_heading(doc, "3. Data source choices", 1)
    add_body(doc, "The workflow choice above is about where each job runs; the choice below is about which inputs to trust. The main revision from the original plan is to make E-OBS the primary temperature source and demote ERA5-Land to a cross-check.")
    add_source_table(doc)

    add_heading(doc, "4. Recommended staged workflow", 1)
    add_stage_table(doc)

    add_heading(doc, "5. Why this sequence is preferred", 1)
    add_body(doc, "Reproducibility comes first. The heatwave definition, population interpolation and raster aggregation should be scripts, so that every annual layer can be regenerated and audited.")
    add_body(doc, "Visual inspection comes second. QGIS is the quickest way to discover shifted grids, incorrect coordinate systems, unexpected no-data and implausible totals before those errors reach an interface.")
    add_body(doc, "Interactivity comes last. Shiny is valuable when the annual layers are stable, but it is a communication surface rather than the place to define or debug the analysis.")

    add_heading(doc, "6. Decision rules for the first prototype", 1)
    rules = [
        ("Use local, compound thresholds", "A heatwave here is a locally relative, day-and-night event: flag a day when both daily maximum and daily minimum temperature exceed their local calendar-day 90th percentile for at least two consecutive days. The two-day minimum follows the health-epidemiology literature; the compound (day+night) criterion is the form most consistently linked to excess mortality (Wang et al. 2025, Nature Communications; intensity via the Excess Heat Factor, Nairn and Fawcett 2015; the P90 / consecutive-day convention as in the EEA Climate-ADAPT health-heatwave indicator)."),
        ("Keep intensity separate", "Retain temperature excess above the thresholds as a separate intensity layer; report the Excess Heat Factor (EHF) as the recognised comparison metric."),
        ("Preserve separate layers", "Keep heatwave days, population and exposed person-days as separate annual rasters, not only as a final composite."),
        ("Conserve population on aggregation", "Move GHS-POP to the 0.1° grid by area-weighted summation (exactextractr with grid cells as polygons), which guarantees count conservation; never use bilinear resampling or a mean, which turn counts into a density."),
        ("Treat age as a filter", "Join 65+ and 75+ country-year shares to exposure summaries after spatial aggregation; do not display an imputed age raster in the first map."),
        ("Document the humidity gap", "The metric uses dry-bulb temperature only. Humidity (E-OBS HU, or apparent temperature) is a known modifier of heat risk; record its omission and keep it as a candidate sensitivity test."),
        ("Use a fixed map scale", "A fixed colour scale across 1991-2020 makes annual changes interpretable; record any values outside the chosen display range."),
    ]
    for label, text in rules:
        add_body(doc, text, bold_label=label + ". ")

    add_heading(doc, "7. Sensitivity and robustness checks", 1)
    add_body(doc, "Define the exposure layer once, then vary one assumption at a time and record how the European total and the country ranking move:")
    checks = [
        ("Climate source", "E-OBS (primary) versus ERA5-Land, especially over coasts and sparse-station regions."),
        ("Event duration", "Two versus three consecutive days above the thresholds."),
        ("Definition family", "The compound TX-and-TN rule versus the WSDI six-day TX-only index."),
        ("Percentile window", "Width of the centred calendar-day window (for example 15 versus 31 days)."),
        ("Humidity", "Dry-bulb temperature versus an apparent-temperature variant using E-OBS relative humidity."),
        ("Population", "Linear GHS-POP interpolation versus WorldPop annual counts for the post-2000 years."),
    ]
    for label, text in checks:
        add_body(doc, text, bold_label=label + ": ")

    add_heading(doc, "8. Practical next steps", 1)
    add_body(doc, "1. Run a small 1991 test: GHS-POP 1990/1995 interpolated, with one month of E-OBS daily TX/TN, and the same month of ERA5-Land for comparison.")
    add_body(doc, "2. Install and test the R spatial stack (terra, exactextractr, sf), confirm count conservation on the test month, then write the reproducible preparation script under renv.")
    add_body(doc, "3. Validate the first outputs in QGIS, including the E-OBS versus ERA5-Land difference, before scaling to the full 1991-2020 series.")
    add_body(doc, "4. Decide on the Shiny interface only after the annual rasters and country-year table are stable.")

    add_heading(doc, "Sources and project context", 1)
    add_body(doc, "This note operationalises the project protocol 'Heatwave population-exposure prototype: protocol' (Phase 3) and revises two of its choices: E-OBS replaces ERA5-Land as the primary temperature source, and the heatwave duration threshold moves to two consecutive days. The protocol should be updated to match. Data sources are E-OBS (CDS insitu-gridded-observations-europe), Copernicus ERA5-Land daily statistics (cross-check), GHS-POP R2023A, GISCO country boundaries and Eurostat demographic tables.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

# =============================================================================
# add_part_iv_to_docx.py
# Inserts Part IV (Multidimensional Exposure Typologies) into the Word Document
# =============================================================================
import docx

doc_path = "Feature explorations/Heatwave exposure mapping methods.docx"
doc = docx.Document(doc_path)

# Find index of References paragraph
ref_p = None
for p in doc.paragraphs:
    if p.text.strip().lower() == "references":
        ref_p = p
        break

if ref_p is None:
    raise ValueError("Could not find 'References' section in document.")

part_iv_paragraphs = [
    ("", "Normal"),
    ("Part IV. Multidimensional Exposure Typologies: Spatial Clustering (K-Means)", "Normal"),
    ("", "Normal"),
    ("16. Rationale and Causal Architecture: Beyond Additive Indices", "Normal"),
    ("A recurring temptation in spatial vulnerability mapping is to collapse heterogeneous indicators into a single composite score—either by simple summation, averaging, or weighted index construction. Under the causal architecture of this project, this approach carries a severe methodological flaw: it smuggles in an implicit assumption of perfect substitutability. An additive index asserts that a higher socio-economic buffer (or lower air pollution) can quantitatively 'cancel out' physiological heat stress or light pollution. In terms of human lived experience, interface features from different spheres (A1 Ecological, A2 Technosphere, A3 Social Organization) act through non-substitutable biological, sensory, and temporal mechanisms.", "Normal"),
    ("To map how the Anthropocene is encountered without flattening these distinct mechanisms, we adopt the methodology established in environmental epidemiology and exposome research: unsupervised spatial clustering (K-Means). Rather than inventing arbitrary weights to aggregate variables into a one-dimensional score, K-Means identifies spatial archetypes—naturally occurring, multi-dimensional profiles of exposure across Europe's geographic space.", "Normal"),
    ("", "Normal"),
    ("17. Variable Selection and Profile Formulations", "Normal"),
    ("Every variable fed into the clustering pipeline is strictly audited to ensure it represents a valid Layer A Experienceable Feature (an interface actually encountered by a human body or daily routine) rather than an upstream Earth-system driver (Step 1 Component), a behavioral/time-use filter (Layer B Filter), or a downstream outcome (Layer C Implication). Three distinct exposure profiles were formulated to capture different axes of Anthropocene experience:", "Normal"),
    ("Profile 1: Sensory & Corporeal Burden. Maps the spatial co-occurrence of direct physical and sensory stressors on the human body, combining Built fraction (A2, WSF3D, % surface), Artificial sky brightness (A2, Falchi et al. 2016, mcd/m²), PM2.5 fine particulate matter (A1, CAMS 2024, µg/m³), and UTCI heat stress hours (A1, ERA5-Land 2022, hours >= 26°C).", "Normal"),
    ("Profile 2: Techno-Ecological Disconnection. Maps the spatial gradient of ecological intactness versus built human mediation, combining Land-use change frequency (A1, HILDA+ 1960–2019), Built fraction (A2, WSF3D), Artificial sky brightness (A2, Falchi et al. 2016), and Population density (A3, GHS-POP, hab/cell).", "Normal"),
    ("Profile 3: Resource & Climate Vulnerability. Maps the compound exposure of atmospheric heat stress, agricultural land occupation, and infrastructure density, combining UTCI heat stress hours (A1, ERA5-Land 2022), Cropland fraction (A3, Potapov et al. 2022, % surface), PM2.5 fine particulate matter (A1, CAMS 2024), and Built fraction (A2, WSF3D).", "Normal"),
    ("", "Normal"),
    ("18. Data Preprocessing, Standardization & K-Means Optimization", "Normal"),
    ("Spatial Harmonization. All datasets are reprojected and resampled to a standardized 3 km European grid under the equal-area LAEA Europe projection (EPSG:3035) with geographic extent [xmin: 2.5e6, xmax: 6.0e6, ymin: 1.5e6, ymax: 5.5e6]. Equal-area projection is mandatory for spatial clustering so that grid cell areas remain uniform across latitudes.", "Normal"),
    ("Distributional Transformations and Z-Score Normalization. Raw variables span incompatible units (hours/year, µg/m³, mcd/m², %). Furthermore, highly skewed distributions (such as built fraction and sky brightness) would distort distance calculations in Euclidean space. We applied a logarithmic transformation y = log(1 + x) to compressed skewed variables, followed by Z-score standardization across all valid European grid cells so that each feature contributes equally (mean = 0, variance = 1) to distance matrix calculations.", "Normal"),
    ("Optimization of Cluster Count (k) via Silhouette Analysis. Rather than imposing an arbitrary number of archetypes, the optimal number of clusters k in [3, 6] is determined empirically using Silhouette score evaluation (Kaufman & Rousseeuw 1990) on a representative spatial sample (N = 10,000 cells). The final k is selected at the global peak of average Silhouette width, yielding k = 3 for Profile 1, k = 4 for Profile 2, and k = 3 for Profile 3.", "Normal"),
    ("", "Normal"),
    ("19. Post-Processing: Spatial Focal Filtering & Accessible Visuals", "Normal"),
    ("Spatial Modal Focal Filter (3x3). Pixel-by-pixel clustering on high-resolution rasters inherently introduces spatial noise ('pepper-and-salt' artifacts) due to local micro-variations. To reveal continuous geographic exposure regions, a 3x3 focal modal filter is applied to the output raster grid. This replaces isolated single-pixel clusters with the predominant surrounding archetype, generating continuous atlas-quality spatial zones.", "Normal"),
    ("Cartographic Standards & Color Safety. In compliance with publication standards (Nature Methods, Okabe & Ito 2008; Paul Tol schemes), maps are built with a soft ocean blue background (#dceeff) to demarcate coastlines, administrative country borders drawn on top of the raster layer in subtle dark grey (grey35) to prevent data occlusion, and discrete colorblind-safe qualitative palettes with explicit descriptive labels.", "Normal"),
    ("", "Normal")
]

for text, style in part_iv_paragraphs:
    p = ref_p.insert_paragraph_before(text)

# Append new references at the end of the document
new_refs = [
    "Falchi, F., Cinzano, P., Duriscoe, D., Kyba, C. C., Elvidge, C. D., Baugh, K., Portnov, B. A., Furgoni, N. A., & Fizek, R. (2016). The new world atlas of artificial night sky brightness. Science Advances, 2(6), e1600377.",
    "Kaufman, L., & Rousseeuw, P. J. (1990). Finding Groups in Data: An Introduction to Cluster Analysis. John Wiley & Sons.",
    "Okabe, M., & Ito, K. (2008). Color Universal Design (CUD) - How to make figures and presentations that are friendly to colorblind people. JFLY Data Vis."
]

for ref in new_refs:
    doc.add_paragraph(ref)

doc.save(doc_path)
print("Successfully inserted Part IV into Heatwave exposure mapping methods.docx!")
